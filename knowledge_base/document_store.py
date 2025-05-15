# knowledge_base/document_store.py
"""
Document processing optimized for fast ingestion into Pinecone.
"""
import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import hashlib

from knowledge_base.schema import Document, DocumentMetadata
from knowledge_base.config import CHUNK_SIZE, CHUNK_OVERLAP

logger = logging.getLogger(__name__)

class DocumentStore:
    """
    Simple document store for processing files into chunks.
    Optimized for speed and Pinecone ingestion.
    """
    
    def __init__(
        self,
        chunk_size: int = CHUNK_SIZE,
        chunk_overlap: int = CHUNK_OVERLAP
    ):
        """Initialize document store."""
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Supported file types
        self.supported_extensions = {
            '.txt', '.md', '.pdf', '.docx', '.html'
        }
        
        logger.info(f"Initialized DocumentStore with chunk_size={chunk_size}")
    
    def load_text(self, text: str, source_name: str = "text_input") -> List[Document]:
        """Load text and split into chunks."""
        # Create metadata
        metadata = DocumentMetadata(
            source=source_name,
            source_type="text",
            chunk_count=0,
            file_size=len(text)
        )
        
        # Split into chunks
        chunks = self._split_text(text)
        
        # Create documents
        documents = []
        for i, chunk in enumerate(chunks):
            doc_metadata = metadata.dict()
            doc_metadata.update({
                "chunk_index": i,
                "chunk_count": len(chunks)
            })
            
            doc = Document(
                text=chunk,
                metadata=doc_metadata,
                doc_id=self._generate_id(source_name, i)
            )
            documents.append(doc)
        
        logger.info(f"Created {len(documents)} chunks from text")
        return documents
    
    def load_file(self, file_path: str) -> List[Document]:
        """Load file and convert to documents."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check file extension
        ext = Path(file_path).suffix.lower()
        if ext not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {ext}")
        
        # Read file content
        try:
            text = self._read_file(file_path)
            source_name = os.path.basename(file_path)
            
            # Create metadata
            file_stat = os.stat(file_path)
            metadata = DocumentMetadata(
                source=source_name,
                source_type="file",
                file_path=file_path,
                file_type=ext,
                file_name=source_name,
                file_size=file_stat.st_size,
                created_at=file_stat.st_ctime,
                modified_at=file_stat.st_mtime
            )
            
            # Split into chunks
            chunks = self._split_text(text)
            
            # Create documents
            documents = []
            for i, chunk in enumerate(chunks):
                doc_metadata = metadata.dict()
                doc_metadata.update({
                    "chunk_index": i,
                    "chunk_count": len(chunks)
                })
                
                doc = Document(
                    text=chunk,
                    metadata=doc_metadata,
                    doc_id=self._generate_id(source_name, i)
                )
                documents.append(doc)
            
            logger.info(f"Loaded {len(documents)} chunks from {file_path}")
            return documents
            
        except Exception as e:
            logger.error(f"Error loading file {file_path}: {e}")
            raise
    
    def load_directory(
        self,
        directory: str,
        recursive: bool = True,
        max_files: Optional[int] = None
    ) -> List[Document]:
        """Load all supported files from directory."""
        if not os.path.exists(directory):
            raise FileNotFoundError(f"Directory not found: {directory}")
        
        # Find files
        files = []
        if recursive:
            for root, _, filenames in os.walk(directory):
                for filename in filenames:
                    file_path = os.path.join(root, filename)
                    if Path(file_path).suffix.lower() in self.supported_extensions:
                        files.append(file_path)
        else:
            for filename in os.listdir(directory):
                file_path = os.path.join(directory, filename)
                if os.path.isfile(file_path) and Path(file_path).suffix.lower() in self.supported_extensions:
                    files.append(file_path)
        
        # Limit files if specified
        if max_files and len(files) > max_files:
            files = files[:max_files]
        
        # Load all files
        all_documents = []
        for file_path in files:
            try:
                documents = self.load_file(file_path)
                all_documents.extend(documents)
            except Exception as e:
                logger.error(f"Error loading {file_path}: {e}")
                continue
        
        logger.info(f"Loaded {len(all_documents)} chunks from {len(files)} files")
        return all_documents
    
    def _read_file(self, file_path: str) -> str:
        """Read file content based on extension."""
        ext = Path(file_path).suffix.lower()
        
        if ext == '.txt' or ext == '.md':
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
        
        elif ext == '.pdf':
            # Simple PDF reading (you might want to use PyMuPDF or similar)
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    text = ""
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
                    return text
            except ImportError:
                logger.warning("PyPDF2 not installed, skipping PDF")
                return ""
        
        elif ext == '.docx':
            # Simple DOCX reading
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                return text
            except ImportError:
                logger.warning("python-docx not installed, skipping DOCX")
                return ""
        
        elif ext == '.html':
            # Simple HTML reading
            try:
                from bs4 import BeautifulSoup
                with open(file_path, 'r', encoding='utf-8') as f:
                    soup = BeautifulSoup(f.read(), 'html.parser')
                    return soup.get_text()
            except ImportError:
                logger.warning("beautifulsoup4 not installed, skipping HTML")
                return ""
        
        else:
            # Fallback to text reading
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
    
    def _split_text(self, text: str) -> List[str]:
        """Split text into overlapping chunks."""
        if not text:
            return []
        
        # Simple word-based chunking
        words = text.split()
        chunks = []
        
        # Calculate words per chunk (approximate)
        words_per_chunk = self.chunk_size // 4  # Rough estimate: 4 chars per word
        overlap_words = self.chunk_overlap // 4
        
        for i in range(0, len(words), words_per_chunk - overlap_words):
            chunk_words = words[i:i + words_per_chunk]
            chunk = " ".join(chunk_words)
            
            # Only add non-empty chunks
            if chunk.strip():
                chunks.append(chunk)
            
            # Break if we've reached the end
            if i + words_per_chunk >= len(words):
                break
        
        return chunks
    
    def _generate_id(self, source_name: str, chunk_index: int) -> str:
        """Generate unique document ID."""
        # Create hash from source name and index
        content = f"{source_name}_{chunk_index}"
        return hashlib.md5(content.encode()).hexdigest()